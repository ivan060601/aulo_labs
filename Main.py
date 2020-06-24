import math
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
import res.math2word as m2w

# Создание документа
doc = Document()

# Задание стиля
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# 1ый параграф для табличек
p1 = doc.add_paragraph('7. Результаты измерений', style='Normal')

# Входные данные для диаметров (м)
inputDataD = [0.38, 0.36, 0.245, 0.25, 0.05, 0.13]

# Таблица с диаметрами
tableD = doc.add_table(rows=2, cols=len(inputDataD) + 1, style='Table Grid')
tableD.alignment = WD_TABLE_ALIGNMENT.CENTER
tableD.cell(1, 0).text = "D"
tableD.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(1, len(inputDataD) + 1):
    tableD.cell(0, i).text = str(i)
    tableD.cell(0, i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    tableD.cell(1, i).text = str(inputDataD[i - 1])
    tableD.cell(1, i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

# Входные данные для высот (м)
inputDataH = [0.24, 0.20, 0.22]

# 2ой параграф, чтобы таблицы не слились
p2 = doc.add_paragraph('', style='Normal')

# Таблица с высотами
tableH = doc.add_table(rows=2, cols=len(inputDataH) + 1, style='Table Grid')
tableH.alignment = WD_TABLE_ALIGNMENT.CENTER
tableH.cell(1, 0).text = "H"
tableH.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(1, len(inputDataH) + 1):
    tableH.cell(0, i).text = str(i)
    tableH.cell(0, i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    tableH.cell(1, i).text = str(inputDataH[i - 1])
    tableH.cell(1, i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

# Коэффициент Стьюдента для 0-10 измерений
t = [0, 0, 12.7, 4.3, 3.18, 2.78, 2.57, 2.45, 2.36, 2.31, 2.26]

# Инструментальная погрешность линейки (м)
inaccuracy_1 = 0.005

# 3ий параграф для обработки измерений
p3 = doc.add_paragraph('\n8. Обработка измерений\n', style='Normal')

### Обработка диаметров
D_avg = sum(inputDataD) / len(inputDataD)
# Запишем это в формулу
d_sum_string = ''
for i in range(0, len(inputDataD) - 1):
    d_sum_string = d_sum_string + str(inputDataD[i]) + ' + '
d_sum_string = d_sum_string + str(inputDataD[len(inputDataD) - 1])

formula1 = '\\overline{D} = \\frac{1}{n}\\sum\\limits_{i=1}^{n} {D_{i}} = \\frac{' + d_sum_string + '}{' + str(
    len(inputDataD)) + '} = %3.3f (м)'
p3._element.append(m2w.math_to_word(formula1 % (D_avg)))

# Вычисление S_D
d_modified = list(map(lambda x: (x - D_avg) ** 2, inputDataD))
S_D = math.sqrt(sum(list(d_modified)) / (len(inputDataD) * (len(inputDataD) - 1)))

# Запишем это в формулу
p3.add_run('\n \n')

d_modified_string = ''
for i in range(0, len(d_modified) - 1):
    d_modified_string = d_modified_string + str(round(d_modified[i], 3)) + ' + '
d_modified_string = d_modified_string + str(round(d_modified[len(d_modified) - 1], 3))

formula2 = 'S_D = \\sqrt {\\frac{\\sum\\limits_{i=1}^n {(D_i-\\overline{D})^2}}{n(n-1)}} = \\sqrt {\\frac{' + d_modified_string + '}{' + str(
    len(inputDataD)) + '(' + str(len(inputDataD)) + '-1)}} = %0.3f'
p3._element.append(m2w.math_to_word(formula2 % (S_D)))

delta_D_avg = S_D * t[len(inputDataD)]
delta_D = math.sqrt(delta_D_avg ** 2 + inaccuracy_1 ** 2)
relative_D = delta_D * 100 / D_avg
# print('D = (%0.3f ± %0.3f) м, ε = %2d, α = 0.95' % (D_avg, delta_D, relative_D))

### Обработка высот
H_avg = sum(inputDataH) / len(inputDataH)
S_H = math.sqrt(sum(list(map(lambda x: (x - H_avg) ** 2, inputDataH))) / (len(inputDataH) * (len(inputDataH) - 1)))
delta_H_avg = S_H * t[len(inputDataH)]
delta_H = math.sqrt(delta_H_avg ** 2 + inaccuracy_1 ** 2)
relative_H = delta_H * 100 / H_avg
# print('H = (%0.3f ± %0.3f) м, ε = %2d, α = 0.95' % (H_avg, delta_H, relative_H))

### Вычисление объема
# V= 1/4 * pi * H * D^2
V_avg = (math.pi * H_avg * D_avg ** 2) / 4
# V'(H) = 1/4 * pi * D^2
vdH = (math.pi * D_avg ** 2) / 4
# V'(D) = 1/2 * pi * H * D
vdD = (math.pi * D_avg * H_avg) / 2
delta_V = math.sqrt((vdH * delta_H) ** 2 + (vdD * delta_D) ** 2)
relative_V = delta_V * 100 / V_avg
# print('V = (%0.3f ± %0.3f) м, ε = %2d, α = 0.95' % (V_avg, delta_V, relative_V))

doc.save('Lab.docx')
